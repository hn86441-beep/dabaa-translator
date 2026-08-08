"""
api/index.py
------------
Backend FastAPI موحّد لتطبيق الترجمة والذكاء الاصطناعي.
مصمم للعمل كدالة Serverless واحدة على Vercel (@vercel/python).

المزوّدون المدعومون:
  - DeepL          -> ترجمة نصية عالية الجودة
  - Groq           -> LLM للترجمة الذكية / التصحيح / المقارنة + Whisper للـ STT + Vision للـ OCR
  - Cohere         -> مزوّد بديل لمهام التصحيح / المقارنة النصية

ملاحظة مهمة حول "المحادثة الجماعية":
  فصل عدة متحدثين من ملف صوتي واحد ممزوج (Speaker Diarization) يتطلب نماذج ML
  ثقيلة (مثل pyannote.audio) لا يمكن تشغيلها ضمن بيئة Vercel Serverless
  (قيود الحجم / الذاكرة / زمن التنفيذ). لذلك تم تصميم هذه الميزة بشكل عملي:
  كل متحدث يُسجّل أو يرفع مقطعه الصوتي الخاص بشكل منفصل، والسيرفر يقوم تلقائياً
  باكتشاف اللغة وترجمة كل مقطع، ثم يجمع كل ذلك في محادثة موحدة مرتبة زمنياً.
"""

import os
import io
import json
import base64
import asyncio
from typing import Optional, List

import httpx
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel

# ---------------------------------------------------------------------------
# إعداد التطبيق
# ---------------------------------------------------------------------------

app = FastAPI(title="AI Translation Suite API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------------------------
# متغيرات البيئة (تُقرأ من Vercel Environment Variables)
# ---------------------------------------------------------------------------

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
DEEPL_API_KEY = os.environ.get("DEEPL_API_KEY", "")
COHERE_API_KEY = os.environ.get("COHERE_API_KEY", "")

# أسماء النماذج قابلة للتعديل عبر متغيرات بيئة اختيارية (بدون تعديل الكود)
GROQ_TEXT_MODEL = os.environ.get("GROQ_TEXT_MODEL", "llama-3.3-70b-versatile")
GROQ_VISION_MODEL = os.environ.get("GROQ_VISION_MODEL", "llama-3.2-11b-vision-preview")
GROQ_STT_MODEL = os.environ.get("GROQ_STT_MODEL", "whisper-large-v3")
COHERE_TEXT_MODEL = os.environ.get("COHERE_TEXT_MODEL", "command-r-plus")

GROQ_BASE_URL = "https://api.groq.com/openai/v1"
COHERE_BASE_URL = "https://api.cohere.com/v1"

HTTP_TIMEOUT = httpx.Timeout(55.0, connect=10.0)  # Vercel Hobby limit ~60s


# ---------------------------------------------------------------------------
# أدوات مساعدة عامة
# ---------------------------------------------------------------------------

def require_key(key: str, name: str):
    if not key:
        raise HTTPException(
            status_code=500,
            detail=f"مفتاح {name} غير موجود في متغيرات البيئة. الرجاء إضافته في إعدادات Vercel.",
        )


def deepl_host() -> str:
    """DeepL يستخدم رابطين مختلفين حسب نوع الحساب (Free تنتهي بـ :fx)."""
    if DEEPL_API_KEY.endswith(":fx"):
        return "https://api-free.deepl.com"
    return "https://api.deepl.com"


async def groq_chat(messages: list, temperature: float = 0.3, max_tokens: int = 2000) -> str:
    require_key(GROQ_API_KEY, "GROQ_API_KEY")
    async with httpx.AsyncClient(timeout=HTTP_TIMEOUT) as client:
        resp = await client.post(
            f"{GROQ_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
            json={
                "model": GROQ_TEXT_MODEL,
                "messages": messages,
                "temperature": temperature,
                "max_tokens": max_tokens,
            },
        )
    if resp.status_code != 200:
        raise HTTPException(status_code=resp.status_code, detail=f"خطأ من Groq: {resp.text}")
    data = resp.json()
    return data["choices"][0]["message"]["content"].strip()


async def cohere_chat(message: str, preamble: Optional[str] = None) -> str:
    require_key(COHERE_API_KEY, "COHERE_API_KEY")
    payload = {"model": COHERE_TEXT_MODEL, "message": message}
    if preamble:
        payload["preamble"] = preamble
    async with httpx.AsyncClient(timeout=HTTP_TIMEOUT) as client:
        resp = await client.post(
            f"{COHERE_BASE_URL}/chat",
            headers={
                "Authorization": f"Bearer {COHERE_API_KEY}",
                "Content-Type": "application/json",
            },
            json=payload,
        )
    if resp.status_code != 200:
        raise HTTPException(status_code=resp.status_code, detail=f"خطأ من Cohere: {resp.text}")
    data = resp.json()
    return data.get("text", "").strip()


# ---------------------------------------------------------------------------
# نماذج البيانات (Pydantic)
# ---------------------------------------------------------------------------

class DeepLTranslateRequest(BaseModel):
    text: str
    target_lang: str            # مثال: "AR", "EN-US", "FR"
    source_lang: Optional[str] = None


class AITranslateRequest(BaseModel):
    text: str
    target_lang: str
    source_lang: Optional[str] = "auto"
    provider: str = "groq"       # "groq" | "cohere"
    task: str = "translate"      # "translate" | "correct" | "compare"
    compare_with: Optional[str] = None   # نص إضافي لمقارنته عند task="compare"
    tone: Optional[str] = "neutral"      # "neutral" | "formal" | "casual"


class DetectLanguageRequest(BaseModel):
    text: str


# ---------------------------------------------------------------------------
# Endpoint: فحص الصحة
# ---------------------------------------------------------------------------

@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "groq_configured": bool(GROQ_API_KEY),
        "deepl_configured": bool(DEEPL_API_KEY),
        "cohere_configured": bool(COHERE_API_KEY),
    }


# ---------------------------------------------------------------------------
# Endpoint: الترجمة عبر DeepL
# ---------------------------------------------------------------------------

@app.post("/api/translate/deepl")
async def translate_deepl(req: DeepLTranslateRequest):
    require_key(DEEPL_API_KEY, "DEEPL_API_KEY")
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="النص فارغ.")

    payload = {
        "text": [req.text],
        "target_lang": req.target_lang.upper(),
    }
    if req.source_lang and req.source_lang.lower() != "auto":
        payload["source_lang"] = req.source_lang.upper()

    async with httpx.AsyncClient(timeout=HTTP_TIMEOUT) as client:
        resp = await client.post(
            f"{deepl_host()}/v2/translate",
            headers={"Authorization": f"DeepL-Auth-Key {DEEPL_API_KEY}"},
            json=payload,
        )
    if resp.status_code != 200:
        raise HTTPException(status_code=resp.status_code, detail=f"خطأ من DeepL: {resp.text}")

    data = resp.json()
    translation = data["translations"][0]
    return {
        "translated_text": translation["text"],
        "detected_source_language": translation.get("detected_source_language"),
        "provider": "deepl",
    }


# ---------------------------------------------------------------------------
# Endpoint: الترجمة / التصحيح / المقارنة الذكية عبر Groq أو Cohere
# ---------------------------------------------------------------------------

@app.post("/api/translate/ai")
async def translate_ai(req: AITranslateRequest):
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="النص فارغ.")

    tone_map = {
        "formal": "استخدم أسلوباً رسمياً واحترافياً.",
        "casual": "استخدم أسلوباً غير رسمي وودوداً.",
        "neutral": "استخدم أسلوباً متوازناً وطبيعياً.",
    }
    tone_instruction = tone_map.get(req.tone, tone_map["neutral"])

    if req.task == "translate":
        system_prompt = (
            "أنت مترجم محترف متعدد اللغات. مهمتك ترجمة النص المُعطى بدقة عالية "
            f"إلى اللغة الهدف: {req.target_lang}. {tone_instruction} "
            "أعد فقط النص المترجم دون أي شرح أو مقدمات."
        )
        user_prompt = req.text

    elif req.task == "correct":
        system_prompt = (
            "أنت مدقق لغوي محترف. صحح الأخطاء الإملائية والنحوية والأسلوبية في النص "
            "التالي مع الحفاظ على المعنى واللغة الأصلية للنص. "
            "أعد فقط النص المصحح دون أي شرح."
        )
        user_prompt = req.text

    elif req.task == "compare":
        if not req.compare_with:
            raise HTTPException(status_code=400, detail="حقل compare_with مطلوب لمهمة المقارنة.")
        system_prompt = (
            "أنت خبير لغوي متخصص في تقييم جودة الترجمة. قارن بين الترجمتين التاليتين "
            "من حيث الدقة والطبيعية والأسلوب، ثم اذكر أيهما أفضل مع تبرير موجز، "
            "وقدم اقتراحاً نهائياً محسّناً إن أمكن. أجب باللغة العربية بشكل منظم."
        )
        user_prompt = (
            f"النص الأصلي:\n{req.text}\n\n"
            f"الترجمة الأولى:\n{req.text}\n\n"
            f"الترجمة الثانية:\n{req.compare_with}"
        )
    else:
        raise HTTPException(status_code=400, detail="task غير مدعومة. استخدم translate/correct/compare.")

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    if req.provider == "groq":
        result = await groq_chat(messages)
    elif req.provider == "cohere":
        result = await cohere_chat(message=user_prompt, preamble=system_prompt)
    else:
        raise HTTPException(status_code=400, detail="provider غير مدعوم. استخدم groq أو cohere.")

    return {"result": result, "provider": req.provider, "task": req.task}


# ---------------------------------------------------------------------------
# Endpoint: اكتشاف اللغة (عبر Groq LLM لضمان دقة عالية بأي لغة)
# ---------------------------------------------------------------------------

@app.post("/api/detect-language")
async def detect_language(req: DetectLanguageRequest):
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="النص فارغ.")
    messages = [
        {
            "role": "system",
            "content": (
                "أعد فقط كود اللغة بصيغة ISO 639-1 (مثل ar, en, fr, es, de, tr) "
                "للنص المُعطى، بدون أي نص إضافي أو شرح."
            ),
        },
        {"role": "user", "content": req.text},
    ]
    lang_code = await groq_chat(messages, temperature=0.0, max_tokens=10)
    lang_code = lang_code.strip().lower().replace(".", "")
    return {"language_code": lang_code}


# ---------------------------------------------------------------------------
# Endpoint: تحويل الصوت إلى نص (STT) عبر Groq Whisper
# ---------------------------------------------------------------------------

@app.post("/api/stt")
async def speech_to_text(
    audio: UploadFile = File(...),
    translate_to: Optional[str] = Form(None),
):
    require_key(GROQ_API_KEY, "GROQ_API_KEY")
    audio_bytes = await audio.read()
    if not audio_bytes:
        raise HTTPException(status_code=400, detail="الملف الصوتي فارغ.")

    files = {"file": (audio.filename or "audio.webm", audio_bytes, audio.content_type or "audio/webm")}
    data = {"model": GROQ_STT_MODEL, "response_format": "verbose_json"}

    async with httpx.AsyncClient(timeout=HTTP_TIMEOUT) as client:
        resp = await client.post(
            f"{GROQ_BASE_URL}/audio/transcriptions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
            data=data,
            files=files,
        )
    if resp.status_code != 200:
        raise HTTPException(status_code=resp.status_code, detail=f"خطأ من Groq STT: {resp.text}")

    result = resp.json()
    transcript = result.get("text", "").strip()
    detected_lang = result.get("language", "unknown")

    response_payload = {
        "transcript": transcript,
        "detected_language": detected_lang,
        "segments": result.get("segments", []),
    }

    if translate_to:
        messages = [
            {
                "role": "system",
                "content": f"ترجم النص التالي بدقة إلى اللغة: {translate_to}. أعد النص المترجم فقط.",
            },
            {"role": "user", "content": transcript},
        ]
        response_payload["translated_text"] = await groq_chat(messages)

    return response_payload


# ---------------------------------------------------------------------------
# Endpoint: استخراج النص من الصور (OCR) عبر نموذج Groq Vision
# ---------------------------------------------------------------------------

@app.post("/api/ocr")
async def ocr_image(
    image: UploadFile = File(...),
    translate_to: Optional[str] = Form(None),
):
    require_key(GROQ_API_KEY, "GROQ_API_KEY")
    image_bytes = await image.read()
    if not image_bytes:
        raise HTTPException(status_code=400, detail="الصورة فارغة.")

    mime = image.content_type or "image/png"
    b64_image = base64.b64encode(image_bytes).decode("utf-8")
    data_url = f"data:{mime};base64,{b64_image}"

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "استخرج كل النص الموجود في هذه الصورة بدقة تامة وبنفس لغته الأصلية، "
                        "دون أي إضافات أو تعليقات أو شرح، فقط النص كما هو."
                    ),
                },
                {"type": "image_url", "image_url": {"url": data_url}},
            ],
        }
    ]

    async with httpx.AsyncClient(timeout=HTTP_TIMEOUT) as client:
        resp = await client.post(
            f"{GROQ_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
            json={
                "model": GROQ_VISION_MODEL,
                "messages": messages,
                "temperature": 0.0,
                "max_tokens": 2000,
            },
        )
    if resp.status_code != 200:
        raise HTTPException(status_code=resp.status_code, detail=f"خطأ من Groq Vision: {resp.text}")

    extracted_text = resp.json()["choices"][0]["message"]["content"].strip()
    response_payload = {"extracted_text": extracted_text}

    if translate_to and extracted_text:
        messages = [
            {
                "role": "system",
                "content": f"ترجم النص التالي بدقة إلى اللغة: {translate_to}. أعد النص المترجم فقط.",
            },
            {"role": "user", "content": extracted_text},
        ]
        response_payload["translated_text"] = await groq_chat(messages)

    return response_payload


# ---------------------------------------------------------------------------
# Endpoint: استخراج النص من الملفات (PDF / DOCX / Excel / TXT)
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")
    return "\n\n".join(pages_text).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"--- {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(c) if c is not None else "" for c in row]
            if any(v.strip() for v in row_values):
                lines.append(" | ".join(row_values))
    return "\n".join(lines).strip()


@app.post("/api/file/extract")
async def extract_file(
    file: UploadFile = File(...),
    translate_to: Optional[str] = Form(None),
):
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="الملف فارغ.")

    filename = (file.filename or "").lower()

    try:
        if filename.endswith(".pdf"):
            extracted_text = extract_text_from_pdf(file_bytes)
        elif filename.endswith(".docx"):
            extracted_text = extract_text_from_docx(file_bytes)
        elif filename.endswith(".xlsx") or filename.endswith(".xlsm"):
            extracted_text = extract_text_from_xlsx(file_bytes)
        elif filename.endswith(".txt") or filename.endswith(".csv"):
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(
                status_code=400,
                detail="صيغة الملف غير مدعومة. الصيغ المدعومة: pdf, docx, xlsx, txt, csv.",
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"تعذّر استخراج النص: {exc}")

    if not extracted_text:
        extracted_text = ""

    response_payload = {"extracted_text": extracted_text, "filename": file.filename}

    # لتفادي تجاوز حدود tokens، نقتصر على أول 12000 حرف عند الترجمة الآلية
    if translate_to and extracted_text:
        text_to_translate = extracted_text[:12000]
        messages = [
            {
                "role": "system",
                "content": f"ترجم النص التالي بدقة إلى اللغة: {translate_to}. أعد النص المترجم فقط.",
            },
            {"role": "user", "content": text_to_translate},
        ]
        response_payload["translated_text"] = await groq_chat(messages, max_tokens=4000)
        if len(extracted_text) > 12000:
            response_payload["truncated"] = True

    return response_payload


# ---------------------------------------------------------------------------
# Endpoint: المحادثة الجماعية (Group Chat)
# ---------------------------------------------------------------------------
#
# آلية العمل:
#   يرسل العميل عدة ملفات صوتية (audios[]) مع أسماء المتحدثين المقابلة (speakers)
#   كسلسلة JSON، مثال: speakers = '["أحمد", "John", "Maria"]'
#   يقوم السيرفر بـ:
#     1. تفريغ كل مقطع صوتي إلى نص (Groq Whisper) + اكتشاف لغته تلقائياً.
#     2. ترجمة كل مقطع إلى اللغة الهدف target_lang.
#     3. إعادة قائمة مرتبة تمثل "المحادثة" مع اسم المتحدث ولغته ونصه الأصلي والمترجم.
# ---------------------------------------------------------------------------

async def transcribe_one(audio_bytes: bytes, filename: str, content_type: str) -> dict:
    files = {"file": (filename, audio_bytes, content_type or "audio/webm")}
    data = {"model": GROQ_STT_MODEL, "response_format": "verbose_json"}
    async with httpx.AsyncClient(timeout=HTTP_TIMEOUT) as client:
        resp = await client.post(
            f"{GROQ_BASE_URL}/audio/transcriptions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
            data=data,
            files=files,
        )
    if resp.status_code != 200:
        raise HTTPException(status_code=resp.status_code, detail=f"خطأ من Groq STT: {resp.text}")
    result = resp.json()
    return {
        "transcript": result.get("text", "").strip(),
        "detected_language": result.get("language", "unknown"),
    }


@app.post("/api/group-chat/process")
async def group_chat_process(
    audios: List[UploadFile] = File(...),
    speakers: str = Form(...),
    target_lang: str = Form("AR"),
):
    require_key(GROQ_API_KEY, "GROQ_API_KEY")
    try:
        speaker_names = json.loads(speakers)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="حقل speakers يجب أن يكون JSON صالحاً (قائمة أسماء).")

    if len(speaker_names) != len(audios):
        raise HTTPException(
            status_code=400,
            detail="عدد أسماء المتحدثين يجب أن يساوي عدد الملفات الصوتية المرفوعة.",
        )

    # تفريغ جميع المقاطع بالتوازي لتقليل زمن الاستجابة (مهم ضمن حدود Vercel الزمنية)
    tasks = []
    for audio in audios:
        audio_bytes = await audio.read()
        tasks.append(transcribe_one(audio_bytes, audio.filename or "audio.webm", audio.content_type))
    transcriptions = await asyncio.gather(*tasks)

    conversation = []
    for speaker_name, transcription in zip(speaker_names, transcriptions):
        transcript = transcription["transcript"]
        detected_lang = transcription["detected_language"]

        translated_text = ""
        if transcript:
            messages = [
                {
                    "role": "system",
                    "content": f"ترجم النص التالي بدقة إلى اللغة: {target_lang}. أعد النص المترجم فقط.",
                },
                {"role": "user", "content": transcript},
            ]
            translated_text = await groq_chat(messages, max_tokens=1000)

        conversation.append({
            "speaker": speaker_name,
            "detected_language": detected_lang,
            "original_text": transcript,
            "translated_text": translated_text,
        })

    return {"target_lang": target_lang, "conversation": conversation}


# ---------------------------------------------------------------------------
# نقطة نهاية جذر بسيطة للتأكد من عمل الـ API (اختياري)
# ---------------------------------------------------------------------------

@app.get("/api")
async def root():
    return {"message": "AI Translation Suite API يعمل بنجاح 🚀"}
